from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


@dataclass
class ParagraphItem:
    text: str
    style: Optional[str] = None
    bold: bool = False
    alignment: Optional[str] = None
    list_type: Optional[str] = None  # "bullet", "number", or None
    font_size: Optional[float] = None  # in points

    def to_plain_text(self) -> str:
        return self.text


@dataclass
class SemanticBlock:
    title: Optional[str] = None
    level: int = 0  # 0 = no heading, 1-3 = Heading 1-3
    items: List[ParagraphItem] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n".join(item.to_plain_text() for item in self.items if item.text)


class DocxParser:
    def __init__(self, path: Path | str) -> None:
        self.path = Path(path)
        self.document = Document(self.path)

    def _is_heading(self, paragraph) -> tuple[bool, int]:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading "):
            try:
                level = int(style_name.split(" ", 1)[1])
                if 1 <= level <= 3:
                    return True, level
            except ValueError:
                pass
        return False, 0

    def _get_font_size(self, paragraph) -> Optional[float]:
        """Get the font size in points from the first run, or None."""
        for run in paragraph.runs:
            if run.font.size is not None:
                return run.font.size.pt
        return None

    def _guess_heading_level(self, paragraph) -> tuple[bool, int]:
        """Heuristic: detect pseudo-headings by font size + bold + alignment."""
        # Already handled by style
        is_heading, level = self._is_heading(paragraph)
        if is_heading:
            return True, level

        text = paragraph.text.strip()
        if not text:
            return False, 0

        font_size = self._get_font_size(paragraph)
        bold = any(run.bold for run in paragraph.runs)
        alignment = paragraph.paragraph_format.alignment

        # Heuristic rules
        # Large font + bold + short text → likely heading
        if font_size and font_size >= 16 and bold and len(text) < 80:
            return True, 1
        if font_size and font_size >= 14 and bold and len(text) < 80:
            return True, 2
        if font_size and font_size >= 12 and bold and alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 80:
            return True, 2

        return False, 0

    def _extract_tables(self) -> List[ParagraphItem]:
        """Extract text from tables as plain text blocks."""
        items = []
        for table in self.document.tables:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows_text.append(" | ".join(cells))
            if rows_text:
                items.append(ParagraphItem(
                    text="\n".join(rows_text),
                    style="Table",
                    bold=False,
                    alignment=None,
                    list_type=None,
                    font_size=None,
                ))
        return items

    def _get_alignment(self, paragraph) -> Optional[str]:
        alignment = paragraph.paragraph_format.alignment
        if alignment is None:
            return None
        # Map enum to string
        mapping = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        # Fallback for older python-docx versions
        try:
            return mapping.get(alignment)
        except Exception:
            return str(alignment)

    def _get_list_type(self, paragraph) -> Optional[str]:
        # Check if paragraph is part of a list via numbering
        pPr = paragraph._element.get_or_add_pPr()
        numPr = pPr.numPr
        if numPr is not None:
            numId = numPr.numId
            if numId is not None:
                # Determine bullet vs number based on abstract number definition
                # Simplified: check if numbering is bullet
                abstract_num_id = self._get_abstract_num_id(numId.val)
                if abstract_num_id is not None:
                    abstract_num = self._get_abstract_num(abstract_num_id)
                    if abstract_num is not None:
                        for lvl in abstract_num.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lvl"):
                            numFmt = lvl.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numFmt")
                            if numFmt is not None:
                                val = numFmt.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
                                if val == "bullet":
                                    return "bullet"
                                else:
                                    return "number"
                return "number"  # default fallback
        return None

    def _get_abstract_num_id(self, num_id_val: int) -> Optional[int]:
        numbering_part = self.document.part.numbering_part
        if numbering_part is None:
            return None
        numbering = numbering_part.numbering_definitions
        for num in numbering.num_lst:
            if num.numId == num_id_val:
                return num.abstract_num_id
        return None

    def _get_abstract_num(self, abstract_num_id: int):
        numbering_part = self.document.part.numbering_part
        if numbering_part is None:
            return None
        numbering = numbering_part.numbering_definitions
        for abstract_num in numbering.abstract_num_lst:
            if abstract_num.abstractNumId == abstract_num_id:
                return abstract_num.element
        return None

    def _extract_paragraph_item(self, paragraph) -> ParagraphItem:
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else None
        bold = any(run.bold for run in paragraph.runs)
        alignment = self._get_alignment(paragraph)
        list_type = self._get_list_type(paragraph)
        font_size = self._get_font_size(paragraph)
        return ParagraphItem(
            text=text,
            style=style,
            bold=bold,
            alignment=alignment,
            list_type=list_type,
            font_size=font_size,
        )

    def _is_in_footer_header(self, paragraph) -> bool:
        """Check if paragraph belongs to header or footer."""
        # Headers and footers are in separate sections; paragraphs in main document.body are not
        # This is a simplistic check; python-docx doesn't expose this directly for paragraphs
        # We rely on the fact that document.paragraphs only yields body paragraphs
        return False

    def parse(self) -> List[SemanticBlock]:
        blocks: List[SemanticBlock] = []
        current_block = SemanticBlock(title=None, level=0)

        for paragraph in self.document.paragraphs:
            is_heading, level = self._guess_heading_level(paragraph)
            if is_heading:
                if current_block.items or current_block.title is not None:
                    blocks.append(current_block)
                current_block = SemanticBlock(
                    title=paragraph.text.strip(),
                    level=level,
                    items=[],
                )
            else:
                item = self._extract_paragraph_item(paragraph)
                if item.text or item.list_type:
                    current_block.items.append(item)

        # Don't forget the last block
        if current_block.items or current_block.title is not None:
            blocks.append(current_block)

        # Merge consecutive no-title blocks if any
        merged: List[SemanticBlock] = []
        for block in blocks:
            if merged and merged[-1].level == 0 and block.level == 0:
                merged[-1].items.extend(block.items)
            else:
                merged.append(block)

        # Append tables as a separate block at the end if any exist
        table_items = self._extract_tables()
        if table_items:
            merged.append(SemanticBlock(
                title="Таблицы",
                level=2,
                items=table_items,
            ))

        return merged


def parse_docx(path: Path | str) -> List[SemanticBlock]:
    parser = DocxParser(path)
    return parser.parse()
